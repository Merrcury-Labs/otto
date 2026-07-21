import json
from datetime import timedelta
from types import SimpleNamespace
from io import BytesIO
from tempfile import TemporaryDirectory
from unittest.mock import Mock, patch

from django.core.files.base import ContentFile
from django.core.cache import cache
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone
from rest_framework.test import APIClient
from docx import Document as DocxDocument
from pydantic import BaseModel

from dashboard.models import Org, Tutor
from users.models import User

from .models import (
    ArtifactReview,
    AIUsageRecord,
    GeneratedArtifact,
    GenerationJob,
    ResearchQuestion,
    ResearchSource,
    SourceChunk,
    SourceDocument,
)
from .services import queue_generation_job
from .tasks import extract_source_document, recover_stalled_jobs, start_generation_job
from .workflow import resume_generation_workflow, run_generation_workflow
from .persistence import persist_approved_course
from .providers import (
    AIBudgetExceeded,
    AIRateLimitExceeded,
    _enforce_rate_limit,
    structured_chat_completion,
)


class TinyStructuredOutput(BaseModel):
    value: str


def fake_web_research_provider(*, job_id=None, questions, brief, documents):
    return {
        'sources': [
            {
                'provider_key': 'official-astronomy-source',
                'type': 'WEB',
                'title': 'Official astronomy reference',
                'url': 'https://example.edu/astronomy',
                'publisher': 'Example University',
                'authors': ['A. Researcher'],
                'published_at': '2026-01-15',
                'reliability_score': 0.9,
            }
        ],
        'findings': [
            {
                'question_query': questions[0]['query'],
                'source_key': 'official-astronomy-source',
                'claim': 'Astronomy studies celestial objects and phenomena beyond Earth.',
                'evidence': 'The reference defines astronomy as the study of celestial objects.',
                'confidence': 0.9,
                'source_locator': {'section': 'Definition'},
            }
        ],
    }


class GenerationJobTestCase(TestCase):
    def setUp(self):
        self.media_directory = TemporaryDirectory()
        self.media_override = override_settings(MEDIA_ROOT=self.media_directory.name)
        self.media_override.enable()
        self.addCleanup(self.media_override.disable)
        self.addCleanup(self.media_directory.cleanup)
        self.organization = Org.objects.create(name='Acme Learning')
        self.tutor = Tutor.objects.create(
            name='Tutor', email='tutor@example.com', org=self.organization
        )
        self.user = User.objects.create(
            name='Requester', email='user@example.com', userID='requester-1'
        )
        self.client = APIClient()

    def job_payload(self):
        return {
            'organization': str(self.organization.id),
            'tutor': str(self.tutor.id),
            'requested_by': str(self.user.id),
            'course_brief': {
                'topic': 'Introduction to astronomy',
                'audience': 'secondary school learners',
            },
        }

    def create_job(self, **overrides):
        values = {
            'organization': self.organization,
            'tutor': self.tutor,
            'requested_by': self.user,
            'course_brief': {'topic': 'Astronomy'},
        }
        values.update(overrides)
        return GenerationJob.objects.create(**values)

    def test_create_job_records_draft_and_event(self):
        response = self.client.post(
            reverse('ai:generation-job-list'), self.job_payload(), format='json'
        )

        self.assertEqual(response.status_code, 201)
        job = GenerationJob.objects.get(pk=response.data['id'])
        self.assertEqual(job.status, GenerationJob.Status.DRAFT)
        self.assertEqual(job.events.get().event_type, 'CREATED')

    def test_tutor_must_belong_to_organization(self):
        other_org = Org.objects.create(name='Other Org')
        payload = self.job_payload()
        payload['organization'] = str(other_org.id)

        response = self.client.post(
            reverse('ai:generation-job-list'), payload, format='json'
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('tutor', response.data)

    def test_queue_is_committed_before_dispatch(self):
        job = self.create_job()
        async_result = Mock(id='celery-task-123')
        dispatcher = Mock(return_value=async_result)

        with self.captureOnCommitCallbacks(execute=True):
            queue_generation_job(job, dispatcher)

        job.refresh_from_db()
        dispatcher.assert_called_once_with(str(job.id))
        self.assertEqual(job.status, GenerationJob.Status.QUEUED)
        self.assertEqual(job.celery_task_id, 'celery-task-123')
        self.assertTrue(job.events.filter(event_type='QUEUED').exists())

    def test_cancel_is_idempotently_rejected_for_terminal_job(self):
        job = self.create_job(status=GenerationJob.Status.COMPLETED)

        response = self.client.post(
            reverse('ai:generation-job-cancel', args=(job.id,)), {}, format='json'
        )

        self.assertEqual(response.status_code, 409)
        job.refresh_from_db()
        self.assertEqual(job.status, GenerationJob.Status.COMPLETED)

    @override_settings(GENERATION_WORKFLOW_RUNNER='')
    def test_worker_fails_clearly_without_langgraph_runner(self):
        job = self.create_job(status=GenerationJob.Status.QUEUED)

        result = start_generation_job.apply(args=(str(job.id),)).get()

        job.refresh_from_db()
        self.assertEqual(result['status'], GenerationJob.Status.FAILED)
        self.assertEqual(job.error_code, 'WORKFLOW_NOT_CONFIGURED')
        self.assertEqual(job.attempt_count, 1)

    @override_settings(
        GENERATION_WORKFLOW_RUNNER='ai.workflow.run_generation_workflow',
        LANGGRAPH_DATABASE_URL='',
    )
    def test_worker_runs_blueprint_graph_to_approval_interrupt(self):
        job = self.create_job(status=GenerationJob.Status.QUEUED)

        result = start_generation_job.apply(args=(str(job.id),)).get()

        job.refresh_from_db()
        self.assertTrue(result['interrupted'])
        self.assertEqual(job.status, GenerationJob.Status.WAITING_FOR_BLUEPRINT_APPROVAL)
        self.assertEqual(job.attempt_count, 1)
        self.assertTrue(job.artifacts.filter(type=GeneratedArtifact.Type.BLUEPRINT).exists())

    @override_settings(GENERATION_JOB_STALE_AFTER_SECONDS=60)
    @patch('ai.tasks.start_generation_job.delay')
    def test_recovery_requeues_stalled_active_job(self, delay):
        job = self.create_job(
            status=GenerationJob.Status.RESEARCHING,
            heartbeat_at=timezone.now() - timedelta(minutes=5),
            attempt_count=1,
            max_attempts=3,
        )

        result = recover_stalled_jobs.run()

        job.refresh_from_db()
        self.assertEqual(result, {'recovered': 1})
        self.assertEqual(job.status, GenerationJob.Status.QUEUED)
        delay.assert_called_once_with(str(job.id))

    @patch('ai.views.extract_source_document.delay')
    def test_upload_document_queues_extraction_after_commit(self, delay):
        job = self.create_job()
        delay.return_value = Mock(id='extract-task-1')
        upload = SimpleUploadedFile(
            'notes.md', b'# Stars\n\nStars produce energy through nuclear fusion.',
            content_type='text/markdown',
        )

        with self.captureOnCommitCallbacks(execute=True):
            response = self.client.post(
                reverse('ai:generation-job-documents', args=(job.id,)),
                {'file': upload},
                format='multipart',
            )

        self.assertEqual(response.status_code, 202)
        document = SourceDocument.objects.get(pk=response.data['id'])
        self.assertEqual(document.status, SourceDocument.Status.UPLOADED)
        self.assertEqual(document.extraction_task_id, 'extract-task-1')
        self.assertEqual(len(document.sha256), 64)
        delay.assert_called_once_with(str(document.id))

    def test_upload_rejects_unsupported_file(self):
        job = self.create_job()
        upload = SimpleUploadedFile(
            'archive.zip', b'not-a-zip', content_type='application/zip'
        )

        response = self.client.post(
            reverse('ai:generation-job-documents', args=(job.id,)),
            {'file': upload},
            format='multipart',
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn('file', response.data)

    def test_extract_text_document_creates_provenance_chunks(self):
        job = self.create_job()
        document = SourceDocument(
            job=job,
            original_filename='lesson.txt',
            content_type='text/plain',
            file_size=80,
            sha256='a' * 64,
        )
        document.file.save(
            'lesson.txt',
            ContentFile(
                b'Gravity attracts objects with mass.\n\nOrbits result from gravity and motion.'
            ),
            save=True,
        )

        result = extract_source_document.apply(args=(str(document.id),)).get()

        document.refresh_from_db()
        self.assertEqual(result['status'], SourceDocument.Status.READY)
        self.assertEqual(document.status, SourceDocument.Status.READY)
        self.assertGreater(document.character_count, 0)
        chunk = document.chunks.get(position=0)
        self.assertIn('Gravity', chunk.content)
        self.assertEqual(len(chunk.content_hash), 64)
        self.assertTrue(job.events.filter(event_type='DOCUMENT_READY').exists())

    def test_extract_docx_preserves_heading(self):
        job = self.create_job()
        payload = BytesIO()
        docx = DocxDocument()
        docx.add_heading('Learning about stars', level=1)
        docx.add_paragraph('A star is a luminous sphere of plasma held together by gravity.')
        docx.save(payload)
        data = payload.getvalue()
        document = SourceDocument(
            job=job,
            original_filename='stars.docx',
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            file_size=len(data),
            sha256='c' * 64,
        )
        document.file.save('stars.docx', ContentFile(data), save=True)

        extract_source_document.apply(args=(str(document.id),)).get()

        chunk = document.chunks.get(position=0)
        self.assertEqual(chunk.heading, 'Learning about stars')
        self.assertIn('luminous sphere', chunk.content)

    def test_queue_rejects_unprocessed_documents(self):
        job = self.create_job()
        SourceDocument.objects.create(
            job=job,
            file='pending.txt',
            original_filename='pending.txt',
            content_type='text/plain',
            file_size=10,
            sha256='b' * 64,
        )

        with self.assertRaisesMessage(ValueError, 'must finish processing'):
            queue_generation_job(job, Mock())

    @override_settings(LANGGRAPH_DATABASE_URL='')
    def test_blueprint_workflow_pauses_with_versioned_artifact(self):
        job = self.create_job(status=GenerationJob.Status.STARTING)

        result = run_generation_workflow(str(job.id))

        job.refresh_from_db()
        artifact = job.artifacts.get(type=GeneratedArtifact.Type.BLUEPRINT, version=1)
        self.assertTrue(result['interrupted'])
        self.assertEqual(job.status, GenerationJob.Status.WAITING_FOR_BLUEPRINT_APPROVAL)
        self.assertEqual(job.progress_percent, 40)
        self.assertEqual(artifact.status, GeneratedArtifact.Status.DRAFT)
        self.assertEqual(artifact.content['objectives'][0]['id'], 'LO-1')
        self.assertTrue(job.graph_checkpoint_id)
        self.assertEqual(job.research_questions.count(), 3)
        self.assertEqual(
            job.research_questions.filter(status=ResearchQuestion.Status.NO_RESULTS).count(),
            3,
        )

    @override_settings(LANGGRAPH_DATABASE_URL='')
    def test_document_research_persists_findings_with_chunk_provenance(self):
        job = self.create_job(status=GenerationJob.Status.STARTING)
        document = SourceDocument.objects.create(
            job=job,
            file='astronomy.txt',
            original_filename='astronomy.txt',
            content_type='text/plain',
            file_size=100,
            sha256='d' * 64,
            status=SourceDocument.Status.READY,
        )
        chunk = SourceChunk.objects.create(
            document=document,
            position=0,
            page_number=2,
            heading='Astronomy foundations',
            content='Astronomy foundational concepts include stars, planets, gravity, and orbits.',
            content_hash='e' * 64,
            character_count=75,
        )

        run_generation_workflow(str(job.id))

        source = job.research_sources.get()
        finding = source.findings.first()
        self.assertEqual(source.type, ResearchSource.Type.DOCUMENT)
        self.assertEqual(finding.source_locator['chunk_id'], str(chunk.id))
        self.assertEqual(finding.source_locator['page_number'], 2)
        self.assertTrue(
            job.research_questions.filter(status=ResearchQuestion.Status.COMPLETED).exists()
        )

        response = self.client.get(reverse('ai:generation-job-research', args=(job.id,)))
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data['sources'][0]['title'], 'astronomy.txt')

    @override_settings(
        LANGGRAPH_DATABASE_URL='',
        COURSE_RESEARCH_PROVIDER='ai.tests.fake_web_research_provider',
    )
    def test_configured_web_provider_records_url_and_quality_metadata(self):
        job = self.create_job(status=GenerationJob.Status.STARTING)

        run_generation_workflow(str(job.id))

        source = job.research_sources.get()
        self.assertEqual(source.type, ResearchSource.Type.WEB)
        self.assertEqual(source.url, 'https://example.edu/astronomy')
        self.assertEqual(source.publisher, 'Example University')
        self.assertEqual(source.reliability_score, 0.9)
        self.assertEqual(source.findings.get().confidence, 0.9)

    @override_settings(
        AI_API_KEY='test-key',
        AI_PROVIDER='test-provider',
        AI_MODEL='test-model',
        AI_INPUT_COST_PER_MILLION='1',
        AI_OUTPUT_COST_PER_MILLION='2',
        AI_MAX_OUTPUT_TOKENS=50,
    )
    @patch('ai.providers._client')
    def test_structured_provider_tracks_tokens_cost_and_schema(self, client_factory):
        job = self.create_job(max_ai_tokens=1000)
        response = SimpleNamespace(
            id='provider-request-1',
            choices=[SimpleNamespace(message=SimpleNamespace(content=json.dumps({'value': 'ok'})))],
            usage=SimpleNamespace(prompt_tokens=10, completion_tokens=5),
        )
        client = client_factory.return_value
        client.chat.completions.create.return_value = response

        result = structured_chat_completion(
            job_id=str(job.id),
            operation='test_schema',
            schema_model=TinyStructuredOutput,
            system_prompt='Return structured test output.',
            payload={'input': 'hello'},
        )

        self.assertEqual(result, {'value': 'ok'})
        usage = job.ai_usage_records.get()
        self.assertEqual(usage.provider, 'test-provider')
        self.assertEqual(usage.total_tokens, 15)
        self.assertEqual(str(usage.estimated_cost_usd), '0.000020')
        request = client.chat.completions.create.call_args.kwargs
        self.assertEqual(request['response_format']['type'], 'json_schema')

        api_response = self.client.get(
            reverse('ai:generation-job-usage', args=(job.id,))
        )
        self.assertEqual(api_response.status_code, 200)
        self.assertEqual(api_response.data['totals']['total_tokens'], 15)

    @override_settings(AI_MAX_OUTPUT_TOKENS=50)
    def test_provider_rejects_request_that_exceeds_job_token_budget(self):
        job = self.create_job(max_ai_tokens=2)

        with self.assertRaises(AIBudgetExceeded):
            structured_chat_completion(
                job_id=str(job.id),
                operation='over_budget',
                schema_model=TinyStructuredOutput,
                system_prompt='A prompt long enough to exceed two approximate tokens.',
                payload={'input': 'hello'},
            )

        self.assertFalse(AIUsageRecord.objects.filter(job=job).exists())

    @override_settings(
        AI_PROVIDER='rate-test-provider',
        AI_MODEL='rate-test-model',
        AI_REQUESTS_PER_MINUTE=1,
    )
    def test_provider_rate_limit_is_enforced(self):
        cache.clear()
        _enforce_rate_limit()
        with self.assertRaises(AIRateLimitExceeded):
            _enforce_rate_limit()

    @override_settings(LANGGRAPH_DATABASE_URL='')
    def test_blueprint_workflow_resumes_after_approval(self):
        job = self.create_job(status=GenerationJob.Status.STARTING)
        run_generation_workflow(str(job.id))

        result = resume_generation_workflow(
            str(job.id), {'decision': 'APPROVE', 'feedback': ''}
        )

        job.refresh_from_db()
        artifact = job.artifacts.get(type=GeneratedArtifact.Type.BLUEPRINT, version=1)
        package = job.artifacts.get(type=GeneratedArtifact.Type.COURSE_PACKAGE, version=1)
        self.assertTrue(result['interrupted'])
        self.assertEqual(job.status, GenerationJob.Status.WAITING_FOR_FINAL_APPROVAL)
        self.assertEqual(artifact.status, GeneratedArtifact.Status.APPROVED)
        self.assertEqual(package.status, GeneratedArtifact.Status.DRAFT)

    @override_settings(LANGGRAPH_DATABASE_URL='')
    def test_final_approval_persists_complete_course_once(self):
        job = self.create_job(status=GenerationJob.Status.STARTING)
        run_generation_workflow(str(job.id))
        resume_generation_workflow(str(job.id), {'decision': 'APPROVE', 'feedback': ''})
        package = job.artifacts.get(type=GeneratedArtifact.Type.COURSE_PACKAGE, version=1)

        result = resume_generation_workflow(
            str(job.id), {'decision': 'APPROVE', 'feedback': ''}
        )

        job.refresh_from_db()
        package.refresh_from_db()
        course = job.result_course
        self.assertFalse(result['interrupted'])
        self.assertEqual(job.status, GenerationJob.Status.COMPLETED)
        self.assertEqual(job.progress_percent, 100)
        self.assertEqual(package.status, GeneratedArtifact.Status.APPROVED)
        self.assertEqual(course.modules.count(), 1)
        self.assertEqual(course.lessons.count(), 1)
        self.assertEqual(course.quizzes.count(), 1)
        self.assertEqual(course.quizzes.get().questions.count(), 5)
        self.assertEqual(course.flashcard_decks.count(), 1)
        self.assertEqual(course.flashcard_decks.get().cards.count(), 10)
        self.assertEqual(course.quizzes.get().status, 'DRAFT')
        self.assertEqual(course.flashcard_decks.get().status, 'DRAFT')

        repeated = persist_approved_course(str(job.id), str(package.id))
        self.assertEqual(repeated.id, course.id)

    @override_settings(LANGGRAPH_DATABASE_URL='')
    def test_blueprint_revision_creates_new_artifact_version(self):
        job = self.create_job(status=GenerationJob.Status.STARTING)
        run_generation_workflow(str(job.id))

        result = resume_generation_workflow(
            str(job.id),
            {'decision': 'REVISE', 'feedback': 'Use more practical examples.'},
        )

        job.refresh_from_db()
        first = job.artifacts.get(version=1)
        second = job.artifacts.get(version=2)
        self.assertTrue(result['interrupted'])
        self.assertEqual(job.status, GenerationJob.Status.WAITING_FOR_BLUEPRINT_APPROVAL)
        self.assertEqual(first.status, GeneratedArtifact.Status.REVISION_REQUESTED)
        self.assertEqual(second.status, GeneratedArtifact.Status.DRAFT)

    @override_settings(LANGGRAPH_DATABASE_URL='')
    @patch('ai.views.resume_generation_job.delay')
    def test_review_endpoint_records_and_dispatches_approval(self, delay):
        job = self.create_job(status=GenerationJob.Status.STARTING)
        run_generation_workflow(str(job.id))
        artifact = job.artifacts.get(version=1)
        delay.return_value = Mock(id='resume-task-1')

        with self.captureOnCommitCallbacks(execute=True):
            response = self.client.post(
                reverse('ai:generation-job-review-blueprint', args=(job.id,)),
                {
                    'decided_by': str(self.user.id),
                    'decision': ArtifactReview.Decision.APPROVE,
                },
                format='json',
            )

        self.assertEqual(response.status_code, 202)
        review = artifact.reviews.get()
        self.assertEqual(review.decided_by, self.user)
        self.assertEqual(review.decision, ArtifactReview.Decision.APPROVE)
        delay.assert_called_once()

    @override_settings(LANGGRAPH_DATABASE_URL='')
    @patch('ai.views.resume_generation_job.delay')
    def test_final_review_endpoint_records_and_dispatches_approval(self, delay):
        job = self.create_job(status=GenerationJob.Status.STARTING)
        run_generation_workflow(str(job.id))
        resume_generation_workflow(str(job.id), {'decision': 'APPROVE', 'feedback': ''})
        package = job.artifacts.get(type=GeneratedArtifact.Type.COURSE_PACKAGE, version=1)
        delay.return_value = Mock(id='final-resume-task-1')

        with self.captureOnCommitCallbacks(execute=True):
            response = self.client.post(
                reverse('ai:generation-job-review-final', args=(job.id,)),
                {
                    'decided_by': str(self.user.id),
                    'decision': ArtifactReview.Decision.APPROVE,
                },
                format='json',
            )

        self.assertEqual(response.status_code, 202)
        self.assertEqual(package.reviews.get().decision, ArtifactReview.Decision.APPROVE)
        delay.assert_called_once()

    @override_settings(LANGGRAPH_DATABASE_URL='')
    def test_final_revision_creates_new_course_package_version(self):
        job = self.create_job(status=GenerationJob.Status.STARTING)
        run_generation_workflow(str(job.id))
        resume_generation_workflow(str(job.id), {'decision': 'APPROVE', 'feedback': ''})

        result = resume_generation_workflow(
            str(job.id),
            {'decision': 'REVISE', 'feedback': 'Add clearer applied examples.'},
        )

        job.refresh_from_db()
        first = job.artifacts.get(type=GeneratedArtifact.Type.COURSE_PACKAGE, version=1)
        second = job.artifacts.get(type=GeneratedArtifact.Type.COURSE_PACKAGE, version=2)
        self.assertTrue(result['interrupted'])
        self.assertEqual(job.status, GenerationJob.Status.WAITING_FOR_FINAL_APPROVAL)
        self.assertEqual(first.status, GeneratedArtifact.Status.REVISION_REQUESTED)
        self.assertEqual(second.status, GeneratedArtifact.Status.DRAFT)

# Create your tests here.
